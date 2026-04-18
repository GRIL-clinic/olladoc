"""
Docx extractor
--------------
Walks a .docx and produces a list of Blocks (Heading, BodyPara,
ListItem, TablePlaceholder, Footnote, Comment) plus the raw table data.
"""

import re
import subprocess
import tempfile
from docx import Document
from lxml import etree
from typing import List

from blocks import (Run, Heading, BodyPara, ListItem, TablePlaceholder,
                    Footnote, Comment, Block)


_HEADING_STYLE_RE = re.compile(r"Heading\s*(\d+)", re.IGNORECASE)
_LIST_STYLE_RE = re.compile(r"List\s+(Number|Bullet|Paragraph)",
                            re.IGNORECASE)
_LIST_MARKER_RE = re.compile(r'^\s*\d+[\.\)]\s+\S')


class DocxExtractor:
    """Reads a .docx and emits Blocks + raw tables."""

    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Regex to match a leading number prefix like "1.", "2.3.", "5.6.1."
    # Allows optional leading whitespace (body headings may be indented).
    _NUM_PREFIX_RE = re.compile(r'^\s*(\d+(?:\.\d+)*\.)\s+')

    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self._num_defs = self._parse_numbering()
        self._num_counters: dict[tuple[str, str], int] = {}
        self._fn_id_map: dict[str, str] = {}  # XML id -> display number
        self._heading_numbers = self._extract_heading_numbers()

    def _extract_heading_numbers(self) -> dict[str, str]:
        """Use LibreOffice to render the docx to plain text and extract
        heading numbers. Returns a map of heading text → number prefix
        (e.g. 'Componentes de la base' → '1.')."""
        heading_map: dict[str, str] = {}
        try:
            with tempfile.TemporaryDirectory() as tmp:
                result = subprocess.run(
                    ["soffice", "--headless", "--convert-to", "txt:Text",
                     str(self.filepath), "--outdir", tmp],
                    capture_output=True, text=True, timeout=30,
                )
                if result.returncode != 0:
                    return heading_map
                # Find the output .txt file
                import glob
                txt_files = glob.glob(f"{tmp}/*.txt")
                if not txt_files:
                    return heading_map
                with open(txt_files[0], "r", encoding="utf-8") as f:
                    for line in f:
                        line = line.strip()
                        m = self._NUM_PREFIX_RE.match(line)
                        if m:
                            prefix = m.group(1)
                            text = line[m.end():].strip()
                            # Strip trailing tab+page number (from TOC lines).
                            text = re.split(r'\t\d*$', text)[0].strip()
                            key = text[:40].strip()
                            if key and key not in heading_map:
                                heading_map[key] = prefix
        except Exception:
            pass
        return heading_map

    def _lookup_heading_number(self, heading_text: str) -> str:
        """Look up the rendered number for a heading by matching its text
        against the LibreOffice-extracted map."""
        key = heading_text[:40].strip()
        return self._heading_numbers.get(key, "")

    def _parse_numbering(self) -> dict:
        """Parse numbering.xml to build a map of
        (numId, ilvl) → {fmt, lvlText, start}."""
        defs = {}
        numbering_part = None
        for rel in self.doc.part.rels.values():
            if "numbering" in rel.reltype:
                numbering_part = rel.target_part
                break
        if numbering_part is None:
            return defs

        root = etree.fromstring(numbering_part.blob)
        W = self._W

        # Map numId -> abstractNumId
        num_to_abstract = {}
        for num in root.findall(f"{{{W}}}num"):
            nid = num.get(f"{{{W}}}numId")
            an = num.find(f"{{{W}}}abstractNumId")
            if nid and an is not None:
                num_to_abstract[nid] = an.get(f"{{{W}}}val")

        # Parse abstract numbering levels
        abstracts: dict[str, dict] = {}
        for an in root.findall(f"{{{W}}}abstractNum"):
            aid = an.get(f"{{{W}}}abstractNumId")
            levels = {}
            for lvl in an.findall(f"{{{W}}}lvl"):
                ilvl = lvl.get(f"{{{W}}}ilvl")
                fmt_el = lvl.find(f"{{{W}}}numFmt")
                txt_el = lvl.find(f"{{{W}}}lvlText")
                start_el = lvl.find(f"{{{W}}}start")
                levels[ilvl] = {
                    "fmt": fmt_el.get(f"{{{W}}}val") if fmt_el is not None else "decimal",
                    "lvlText": txt_el.get(f"{{{W}}}val") if txt_el is not None else "%1.",
                    "start": int(start_el.get(f"{{{W}}}val")) if start_el is not None else 1,
                }
            abstracts[aid] = levels

        # Flatten: (numId, ilvl) -> level def
        for nid, aid in num_to_abstract.items():
            if aid in abstracts:
                for ilvl, lvl_def in abstracts[aid].items():
                    defs[(nid, ilvl)] = lvl_def
        return defs

    def _resolve_numbering(self, para) -> str:
        """If a paragraph has w:numPr, resolve and return its display
        number (e.g. '1.' or '1.2.'). Handles multi-level numbering
        with sub-level counter resets."""
        W = self._W
        ppr = para._p.find(f"{{{W}}}pPr")
        if ppr is None:
            return ""
        numPr = ppr.find(f"{{{W}}}numPr")
        if numPr is None:
            return ""
        ilvl_el = numPr.find(f"{{{W}}}ilvl")
        numId_el = numPr.find(f"{{{W}}}numId")
        if ilvl_el is None or numId_el is None:
            return ""
        ilvl = int(ilvl_el.get(f"{{{W}}}val", "0"))
        numId = numId_el.get(f"{{{W}}}val", "0")

        key = (numId, str(ilvl))
        lvl_def = self._num_defs.get(key)
        if not lvl_def:
            return ""

        # Increment counter for this level.
        counter = self._num_counters.get(key, lvl_def["start"] - 1) + 1
        self._num_counters[key] = counter

        # Reset all sub-level counters (they restart when parent increments).
        for sub_ilvl in range(ilvl + 1, 10):
            sub_key = (numId, str(sub_ilvl))
            if sub_key in self._num_counters:
                sub_def = self._num_defs.get(sub_key, {})
                self._num_counters[sub_key] = sub_def.get("start", 1) - 1

        # Build display text: replace %1, %2, %3... with counters
        # from each level.
        result = lvl_def["lvlText"]
        for lvl_idx in range(ilvl + 1):
            lvl_key = (numId, str(lvl_idx))
            lvl_counter = self._num_counters.get(lvl_key, 0)
            result = result.replace(f"%{lvl_idx + 1}", str(lvl_counter))
        return result

    # ---- helpers --------------------------------------------------------

    @staticmethod
    def _heading_level(para) -> int:
        try:
            name = (para.style.name or "").strip()
        except Exception:
            return 0
        m = _HEADING_STYLE_RE.match(name)
        return max(1, min(6, int(m.group(1)))) if m else 0

    @staticmethod
    def _is_list_paragraph(para) -> bool:
        try:
            name = (para.style.name or "").strip()
        except Exception:
            name = ""
        if _LIST_STYLE_RE.search(name):
            return True
        # Numbered/bulleted via numPr element
        ppr = para._p.find(f"{{{DocxExtractor._W}}}pPr")
        if ppr is not None and ppr.find(f"{{{DocxExtractor._W}}}numPr") is not None:
            return True
        return _LIST_MARKER_RE.match(para.text or "") is not None

    def _para_runs(self, para) -> List[Run]:
        """Convert a paragraph's runs into Run model. Footnote
        references (inline superscripts with no <w:t>) are emitted as
        sentinel markers of the form ``‹FN{id}›`` so they survive
        translation and can be re-rendered as superscript runs by the
        translator's docx writer."""
        out: List[Run] = []
        fn_tag = f"{{{self._W}}}footnoteReference"
        for r in para.runs:
            text = r.text or ""
            fn_ref = r._r.find(fn_tag)
            if fn_ref is not None:
                raw_id = fn_ref.get(f"{{{self._W}}}id") or ""
                fn_id = self._fn_id_map.get(raw_id, raw_id)
                text += f"‹FN{fn_id}›"
            if not text:
                continue
            out.append(Run(text=text,
                           bold=bool(r.bold),
                           italic=bool(r.italic),
                           size=float(r.font.size.pt) if r.font.size else 0.0))
        return Run.consolidate(out)

    # ---- table extraction ----------------------------------------------

    @staticmethod
    def _extract_table_rows(table):
        """Extract table cell text, deduping merged cells in each row."""
        rows = []
        for row in table.rows:
            deduped = []
            for cell in row.cells:
                if not deduped or cell.text != deduped[-1]:
                    deduped.append(cell.text)
            rows.append(deduped)
        return rows

    def _extract_table_rows_xml(self, tbl_elem):
        """Extract table rows directly from a w:tbl XML element.
        Used for tables wrapped in SDTs that python-docx doesn't expose."""
        W = self._W
        rows = []
        for tr in tbl_elem.findall(f"{{{W}}}tr"):
            cells = []
            for tc in tr.findall(f"{{{W}}}tc"):
                text = "".join(
                    t.text for t in tc.iter(f"{{{W}}}t") if t.text)
                if not cells or text != cells[-1]:
                    cells.append(text)
            rows.append(cells)
        return rows

    # ---- footnote / comment extraction (docx XML) ---------------------

    def _extract_footnotes(self) -> List[Footnote]:
        out: List[Footnote] = []
        for rel in self.doc.part.rels.values():
            if "footnote" not in rel.reltype.lower():
                continue
            root = etree.fromstring(rel.target_part.blob)
            for fn in root.findall(f"{{{self._W}}}footnote"):
                fn_type = fn.get(f"{{{self._W}}}type", "")
                if fn_type in ("separator", "continuationSeparator"):
                    continue
                xml_id = fn.get(f"{{{self._W}}}id", "")
                texts = [t.text for t in fn.iter(f"{{{self._W}}}t") if t.text]
                text = "".join(texts).strip()
                if text:
                    display = str(len(out) + 1)
                    self._fn_id_map[xml_id] = display
                    out.append(Footnote(marker=display, text=text))
        return out

    @staticmethod
    def _comment_text(comment_elem, ns):
        return "".join(t.text for t in comment_elem.iter(f"{{{ns}}}t")
                       if t.text).strip()

    @classmethod
    def _anchor_text(cls, body, comment_id):
        ns = cls._W
        start = body.find(f".//{{{ns}}}commentRangeStart[@{{{ns}}}id='{comment_id}']")
        end = body.find(f".//{{{ns}}}commentRangeEnd[@{{{ns}}}id='{comment_id}']")
        if start is None or end is None:
            return ""
        collecting = False
        texts = []
        for elem in body.iter():
            if elem is start:
                collecting = True
                continue
            if elem is end:
                break
            if collecting and elem.tag == f"{{{ns}}}t" and elem.text:
                texts.append(elem.text)
        return "".join(texts).strip()

    def _extract_comments(self) -> List[Comment]:
        """Pull review comments from docx comments, ordered by
        their anchor position in the body."""
        comments_part = None
        for rel in self.doc.part.rels.values():
            if rel.reltype.endswith("/comments"):
                comments_part = rel.target_part
                break
        if comments_part is None:
            return []

        body = self.doc.element.body
        body_order = {}
        for i, elem in enumerate(body.iter()):
            if elem.tag == f"{{{self._W}}}commentRangeStart":
                body_order[elem.get(f"{{{self._W}}}id")] = i

        out: List[Comment] = []
        for elem in comments_part._element.findall(f"{{{self._W}}}comment"):
            cid = elem.get(f"{{{self._W}}}id")
            text = self._comment_text(elem, self._W)
            if not text:
                continue
            out.append(Comment(
                id=cid,
                author=elem.get(f"{{{self._W}}}author", ""),
                date=elem.get(f"{{{self._W}}}date", ""),
                text=text,
                anchor=self._anchor_text(body, cid),
            ))
        out.sort(key=lambda c: body_order.get(c.id, float("inf")))
        return out

    # ---- main extract --------------------------------------------------

    def extract(self) -> tuple[List[Block], list]:
        body = self.doc.element.body
        # Map xml elements to python-docx Paragraph / Table objects.
        para_by_elem = {p._p: p for p in self.doc.paragraphs}
        table_by_elem = {t._tbl: t for t in self.doc.tables}

        # Extract footnotes first so _fn_id_map is populated before
        # _para_runs encounters inline footnote references.
        footnotes = self._extract_footnotes()

        blocks: List[Block] = []
        tables: list = []
        table_idx = 0

        # Iterate body children, unwrapping structured document tags
        # (w:sdt) which can contain tables and paragraphs.
        def _body_children(parent):
            for child in parent:
                if child.tag == f"{{{self._W}}}sdt":
                    content = child.find(f"{{{self._W}}}sdtContent")
                    if content is not None:
                        yield from content
                else:
                    yield child

        for child in _body_children(body):
            tag = child.tag
            if tag == f"{{{self._W}}}tbl":
                tbl = table_by_elem.get(child)
                if tbl is not None:
                    tables.append(self._extract_table_rows(tbl))
                else:
                    # Table inside an SDT or other wrapper — not in
                    # doc.tables, so extract directly from XML.
                    tables.append(self._extract_table_rows_xml(child))
                table_idx += 1
                blocks.append(TablePlaceholder(index=table_idx))
                continue
            if tag != f"{{{self._W}}}p":
                continue
            para = para_by_elem.get(child)
            if para is None or not (para.text or "").strip():
                continue

            level = self._heading_level(para)
            runs = self._para_runs(para)
            if not runs:
                continue

            if level > 0:
                full_text = "".join(r.text for r in runs)
                num_prefix = self._lookup_heading_number(full_text)
                if num_prefix:
                    runs[0] = Run(
                        text=num_prefix + " " + runs[0].text,
                        bold=runs[0].bold,
                        italic=runs[0].italic,
                        size=runs[0].size,
                    )
                blocks.append(Heading(level=level, runs=runs))
                continue

            if self._is_list_paragraph(para):
                # Only treat as ListItem if the text starts with an
                # explicit marker (e.g. "a)", "1)", "2."). List Paragraph
                # styles without an explicit marker are just regular
                # paragraphs — treating them as ListItem would split the
                # first word off as an italic "title", causing the first
                # word to be duplicated in the output.
                text = para.text or ""
                if re.match(r'^\s*(?:[a-zA-Z]\)|\d+[\.\)])\s+\S', text):
                    title_text = text.split(None, 1)[0]
                    blocks.append(ListItem(
                        title=title_text, body_runs=runs, separator=" "))
                    continue
                # Auto-numbered List Paragraph (numbering rendered by
                # Word): try to resolve the number prefix and prepend it.
                num_prefix = self._lookup_heading_number(
                    "".join(r.text for r in runs))
                if num_prefix:
                    runs[0] = Run(
                        text=num_prefix + " " + runs[0].text,
                        bold=runs[0].bold,
                        italic=runs[0].italic,
                        size=runs[0].size,
                    )
                    # If the paragraph is all-bold + auto-numbered, it's
                    # visually a heading. Promote it so the renderer
                    # treats it as one. Depth of the number dictates level.
                    all_bold = all(r.bold for r in runs if r.stripped_len)
                    if all_bold:
                        depth = num_prefix.count(".")
                        level = max(2, min(4, depth + 1))
                        blocks.append(Heading(level=level, runs=runs))
                        continue

            blocks.append(BodyPara(runs=runs))

        # Footnotes and comments are appended after body blocks; the
        # translator routes them to their own output files, so
        # their position in the list is incidental.
        blocks.extend(footnotes)
        blocks.extend(self._extract_comments())

        return blocks, tables

    def close(self):
        pass  # python-docx doesn't hold a file handle once Document() returns
