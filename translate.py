"""
Translation Classes
-------------------
Translator: Translates text using a local LLM via Ollama.
TableTranslator: Translates tables and renders them to .docx.
FootnoteTranslator: Translates footnotes and renders them to .docx.
CommentTranslator: Translates comments and renders them to .docx.
DocumentTranslator: Translates a Block list and renders it to .docx.
    Format-agnostic — used by both the PDF and docx flows.
"""

import io
import json
import re
import shutil
import ollama

from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor

from blocks import (Run, Heading, BodyPara, ListItem, Footnote, Comment,
                    ImageBlock, TablePlaceholder, Separator, Block)
from glossary import DomainGlossary, GlossaryEntry
from prompts import TRANSLATEGEMMA_PROMPT, DEFAULT_PROMPT, DEFAULT_DOMAIN


_BLOCK_IDX_RE = re.compile(r'^\[(\d+)\]\s*')
_LIST_MARKER_RE = re.compile(r'^\s*(\d+[\).])\s*')
_FN_MARKER_RE = re.compile(r'‹FN(\d+)›')
_SENT_SPLIT_RE = re.compile(
    r'(?<=[.!?][""\u201d\u2019\')])\s+'   # sentence end after closing quote
    r'|(?<=[.!?])\s+'                      # plain sentence end
)

# Matches markdown emphasis spans and ‹FN› footnote refs in one pass.
# Groups: 1=bold-italic text, 2=bold text, 3=italic text, 4=FN id
_SPAN_RE = re.compile(
    r'\*\*\*(.+?)\*\*\*'   # ***bold-italic***
    r'|\*\*(.+?)\*\*'        # **bold**
    r'|\*(.+?)\*'              # *italic*
    r'|‹FN(\d+)›'     # footnote ref
)

# Path to a log file that accumulates translation warnings (hallucinations,
# prompt echoes, etc.). Callers can set this before starting a translation.
WARNINGS_LOG_PATH: str | None = None


def _log_warning(msg: str) -> None:
    """Print a warning to stdout and, if WARNINGS_LOG_PATH is set, also append it to that file."""
    print(msg)
    if WARNINGS_LOG_PATH:
        try:
            from datetime import datetime
            with open(WARNINGS_LOG_PATH, "a", encoding="utf-8") as f:
                ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                f.write(f"[{ts}] {msg.strip()}\n")
        except Exception:
            pass


def _num_predict_cap(text_len: int) -> int:
    """Cap output tokens proportional to source length so the model can't run off into hallucinated commentary on short inputs.
    Rough budget: ~0.33 tokens/src_char typical output, doubled for safety, with a floor for very short fragments."""
    return max(30, text_len * 2 // 3)


def _translate_preserving_fn(text: str, translate_fn) -> str:
    """Translate text containing ‹FN{id}› markers.

    Strips markers, translates the full text as one unit (preserving context and quality), 
    then re-inserts each marker at the end of its corresponding sentence in the translation.
    """
    fn_matches = list(_FN_MARKER_RE.finditer(text))
    if not fn_matches:
        return translate_fn(text)

    # Figure out which source sentence each FN belongs to.
    clean_src = _FN_MARKER_RE.sub("", text)
    src_sents = _SENT_SPLIT_RE.split(clean_src)
    src_sents = [s for s in src_sents if s.strip()]

    fn_map: list[tuple[int, str]] = []  # (source_sent_idx, fn_id)
    for m in fn_matches:
        offset = len(_FN_MARKER_RE.sub("", text[:m.start()]))
        cumulative = 0
        sent_idx = len(src_sents) - 1
        for si, sent in enumerate(src_sents):
            cumulative += len(sent) + 1
            if offset < cumulative:
                sent_idx = si
                break
        fn_map.append((sent_idx, m.group(1)))

    # Translate the full clean text.
    translated = translate_fn(clean_src)

    # Split translation into sentences.
    trans_sents = _SENT_SPLIT_RE.split(translated)
    trans_sents = [s for s in trans_sents if s.strip()]

    # Attach each FN to the corresponding sentence in the translation.
    # If translation has fewer sentences, clamp to the last one.
    for sent_idx, fn_id in fn_map:
        idx = min(sent_idx, len(trans_sents) - 1)
        trans_sents[idx] = trans_sents[idx].rstrip() + f"‹FN{fn_id}›"

    return " ".join(trans_sents)


class Translator:
    """Translates text using a local LLM via Ollama."""

    # Codes for every language the app offers (TranslateGemma's evaluated set). Keep in sync with LANGS in static/app.js.
    # Per the official spec: bare ISO 639-1 codes, or regionalized language-country pairs. Chinese and Filipino need the regionalized form since they have no plain 639-1 fit (WMT24++ evaluated them as zh_CN and fil_PH).
    GEMMA_LANG_CODES = {
        "Arabic": "ar", "Bengali": "bn", "Bulgarian": "bg", "Catalan": "ca",
        "Chinese": "zh-CN", "Croatian": "hr", "Czech": "cs", "Danish": "da",
        "Dutch": "nl", "English": "en", "Estonian": "et", "Filipino": "fil-PH",
        "Finnish": "fi", "French": "fr", "German": "de", "Greek": "el",
        "Gujarati": "gu", "Hebrew": "he", "Hindi": "hi", "Hungarian": "hu",
        "Icelandic": "is", "Indonesian": "id", "Italian": "it",
        "Japanese": "ja", "Kannada": "kn", "Korean": "ko", "Latvian": "lv",
        "Lithuanian": "lt", "Malayalam": "ml", "Marathi": "mr",
        "Norwegian": "no", "Persian": "fa", "Polish": "pl",
        "Portuguese": "pt", "Punjabi": "pa", "Romanian": "ro",
        "Russian": "ru", "Serbian": "sr", "Slovak": "sk", "Slovenian": "sl",
        "Spanish": "es", "Swahili": "sw", "Swedish": "sv", "Tamil": "ta",
        "Telugu": "te", "Thai": "th", "Turkish": "tr", "Ukrainian": "uk",
        "Urdu": "ur", "Vietnamese": "vi", "Zulu": "zu",
    }

    # Matches snake_case identifiers: two or more lowercase/digit words joined by underscores
    _VAR_RE = re.compile(r'\b([a-z][a-z0-9]*(?:_[a-z0-9]+)+)\b')
    _VAR_PLACEHOLDER = "⟪V{n}⟫"
    _VAR_PLACEHOLDER_RE = re.compile(r'⟪V(\d+)⟫')

    def __init__(self, source_lang, target_lang="English", model="translategemma",
                 model_temp=0.3, glossary=None, verbose_glossary=False,
                 seed=42, domain=DEFAULT_DOMAIN):
        self.source_lang = source_lang
        self.target_lang = target_lang
        # Subject-matter persona injected into the translation prompt. Empty string = general translator, no specialization clause.
        self.domain = (domain or "").strip()
        # A custom domain creates a new echo-able phrase; watch for it alongside the static _PROMPT_LEAK set.
        self._domain_leak = {f"specializing in {self.domain}"} if self.domain else set()
        # Default model is TranslateGemma:
        # https://blog.google/innovation-and-ai/technology/developers-tools/translategemma/
        self.model = model
        self.model_temp = model_temp  # Low temp for more faithful translations
        # When True, prints the glossary entries injected into each chunk's prompt
        self.verbose_glossary = verbose_glossary
        # When set, passed to ollama as `options.seed` for reproducibility.
        self.seed = seed

        # Accept DomainGlossary, plain dict, or None.
        if isinstance(glossary, DomainGlossary):
            self.glossary = glossary
        elif isinstance(glossary, dict) and glossary:
            self.glossary = DomainGlossary.from_dict(glossary)
        else:
            self.glossary = None

    # ---- variable masking --------------------------------------------------

    @classmethod
    def _mask_variables(cls, text: str) -> tuple[str, dict[int, str]]:
        """Replace snake_case identifiers with ⟪Vn⟫ placeholders.

        Returns (masked_text, {n: original_var}).
        """
        registry: dict[int, str] = {}

        def replace(m):
            n = len(registry)
            registry[n] = m.group(1)
            return cls._VAR_PLACEHOLDER.format(n=n)

        return cls._VAR_RE.sub(replace, text), registry

    @classmethod
    def _restore_variables(cls, text: str, registry: dict[int, str]) -> str:
        """Restore ⟪Vn⟫ placeholders to their original identifiers."""
        def replace(m):
            n = int(m.group(1))
            return registry.get(n, m.group(0))
        return cls._VAR_PLACEHOLDER_RE.sub(replace, text)

    # ---- prompt building ---------------------------------------------------

    def _build_prompt(self, text: str, violation_hint: str = "",
                      glossary_section: str | None = None) -> str:
        if glossary_section is None:
            glossary_section = (
                self.glossary.prompt_section(text) if self.glossary else ""
            )
        if self.verbose_glossary and glossary_section:
            preview = text.strip().replace("\n", " ")[:80]
            print(f"  [glossary] chunk {preview!r}:\n"
                  + "\n".join(f"    {l}" for l in glossary_section.splitlines() if l.strip()))
        extra = f"\n{violation_hint}\n" if violation_hint else ""
        if "translategemma" in self.model:
            base = TRANSLATEGEMMA_PROMPT.format(
                source_lang=self.source_lang,
                target_lang=self.target_lang,
                src_code=self.GEMMA_LANG_CODES.get(self.source_lang, "es"),
                tgt_code=self.GEMMA_LANG_CODES.get(self.target_lang, "en"),
                specialization=f", specializing in {self.domain}" if self.domain else "",
                glossary_section=glossary_section,
                text=text,
            )
        else:
            base = DEFAULT_PROMPT.format(
                source_lang=self.source_lang,
                target_lang=self.target_lang,
                specialization=f" for {self.domain}" if self.domain else "",
                glossary_section=glossary_section,
                text=text,
            )
        return base + extra

    def prompt_preview(self) -> str:
        """The Phase 2 prompt with placeholders for chunk text and glossary entries. Shown in logs and the UI so persona, model, and language changes are visible without running a translation."""
        return self._build_prompt(
            "[the chunk of document text being translated goes here]",
            glossary_section="[glossary entries relevant to this chunk go here]\n",
        )

    # ---- output cleaning ---------------------------------------------------

    _PROMPT_LEAK = {"Use established legal phrasing",
                    "Produce ONLY the", "Use standard domain terminology",
                    "Output ONLY the",
                    # Persona and rule phrasing from TRANSLATEGEMMA_PROMPT / DEFAULT_PROMPT
                    "translator working from",
                    "specializing in human rights and public law",
                    "Produce fluent, idiomatic",
                    "Do not translate word-for-word",
                    "Preserve inline markdown emphasis",
                    "The glossary tells you WHAT to say",
                    # Glossary block header from DomainGlossary.prompt_section
                    "use these translations exactly",
                    # Retry-hint header from DomainGlossary.retry_hint_with_previous
                    "The previous translation had terminology errors",
                    "Correct only these specific terms",
                    # Retry-hint header from DomainGlossary.retry_hint_minimal
                    "For this translation, use exactly these terms",
                    # Retry-hint scaffolding from variant B
                    "You previously translated this passage as"}

    # Matches the per-violation bullet lines from retry_hint:
    #   "  - 'foo' must translate to bar"
    _HINT_BULLET_RE = re.compile(r"['\"][^'\"]+['\"] must translate to ")

    def _is_prompt_echo(self, result):
        # Case-insensitive: leaks often come back re-cased (e.g. as a Title-Case heading).
        lowered = result.casefold()
        phrases = self._PROMPT_LEAK | getattr(self, "_domain_leak", set())
        if any(phrase.casefold() in lowered for phrase in phrases):
            return True
        # Also catch partial echoes of the violation hint bullets — the model
        # sometimes copies the "'X' must translate to Y" lines into its output.
        return bool(self._HINT_BULLET_RE.search(result))

    # ---- core translate call -----------------------------------------------

    def _call_model(self, prompt: str, num_predict: int,
                    temp: float | None = None) -> str:
        options = {
            "temperature": temp if temp is not None else self.model_temp,
            "num_predict": num_predict,
        }
        if self.seed is not None:
            options["seed"] = self.seed
        resp = ollama.chat(
            model=self.model,
            messages=[{"role": "user", "content": prompt}],
            options=options,
        )
        return resp["message"]["content"].strip()

    def translate(self, text: str) -> str:
        masked_text, var_registry = self._mask_variables(text)
        prompt = self._build_prompt(masked_text)
        num_predict = _num_predict_cap(len(masked_text))

        result = self._call_model(prompt, num_predict)

        # Prompt-echo guard (retry at low temp).
        if self._is_prompt_echo(result):
            result = self._call_model(prompt, num_predict, temp=0.1)
        if self._is_prompt_echo(result):
            _log_warning(f"  Warning: prompt echo detected, keeping original: {text[:80]!r}")
            return text


        # Glossary violation guard (retry with explicit correction hint).
        if self.glossary:
            violations = self.glossary.violations(masked_text, result)
            if violations:
                # Variant B (default): show the model its previous attempt and ask it to correct specific terms.
                # Swap to `self.glossary.retry_hint_minimal(violations)` for variant A (no previous translation; re-translate cold).
                hint = self.glossary.retry_hint_with_previous(violations, result)
                _log_warning(
                    f"  Glossary violation(s) — retrying with correction hint:\n"
                    + "\n".join(f"    • {v}" for v in violations)
                )
                retry_prompt = self._build_prompt(masked_text, violation_hint=hint)
                retry = self._call_model(retry_prompt, num_predict, temp=0.1)
                if (not self._is_prompt_echo(retry) and
                        len(self.glossary.violations(masked_text, retry)) < len(violations)):
                    result = retry
                else:
                    _log_warning("  Retry did not improve violations; keeping first result.")

        return self._restore_variables(result, var_registry)


class TableTranslator:
    """Translates table cell data and renders the result to .docx."""

    # Characters stripped before checking what's left. 
    # If only digits, whitespace, currency symbols, common punctuation, ranges, or unit separators remain, 
    # the cell is structural data that should be pass through.
    _DATA_STRIP_RE = re.compile(r'[\d\s.,\-/+:%·–—$€£¥@#()\[\]]+')

    @classmethod
    def _is_pure_data_cell(cls, cell: str) -> bool:
        """True if the cell is essentially numeric/code data.

        Strips digits, whitespace, currency symbols, common punctuation, and date/range separators. 
        If <=3 alphabetic chars remain, treat as data and skip translation (avoids LLM hallucinating around bare numbers).

        Examples that match (pass-through):
          "42", "$100", "2023", "01/15/2024", "12 kg", "CAT 1", "v2.1"
        """
        s = cell.strip()
        if not s:
            return False
        stripped = cls._DATA_STRIP_RE.sub('', s)
        alpha_only = ''.join(c for c in stripped if c.isalpha())
        return len(alpha_only) <= 3

    def __init__(self, translator):
        self.translator = translator

    def translate(self, tables):
        # Cache translations of identical cells across all tables.
        translated = []
        cache: dict[str, str] = {}
        for table in tables:
            translated_rows = []
            for row in table:
                new_row = []
                for cell in row:
                    if not cell.strip():
                        new_row.append(cell)
                        continue
                    # Skip cells that are essentially numbers / codes / dates
                    if self._is_pure_data_cell(cell):
                        new_row.append(cell)
                        continue
                    key = cell.strip()
                    if key not in cache:
                        cache[key] = self._translate_cell(cell)
                    new_row.append(cache[key])
                translated_rows.append(new_row)
            translated.append(translated_rows)
        return translated

    # Phrases that, if they appear in a retry's output, mean the model echoed the retry instructions instead of translating.
    _RETRY_ECHO_PHRASES = (
        "This is a short fragment",
        "table cell",
        "Translate ONLY this text",
        "Do not add any explanation",
    )

    def _translate_cell(self, cell: str, max_ratio: float = 4.0) -> str:
        """Translate a table cell.
        If the result is drastically longer than the source (LLM hallucination from a short input), retry once at low temperature. 
        If still bad, fall back to the original text."""

        def _looks_hallucinated(text):
            return text and len(text) > len(cell) * max_ratio + 50

        def _is_retry_echo(text):
            return text and any(p in text for p in self._RETRY_ECHO_PHRASES)

        result = self.translator.translate(cell)
        if _looks_hallucinated(result):
            _log_warning(f"  Warning: table cell translation looks "
                         f"hallucinated ({len(cell)} chars → {len(result)} "
                         f"chars); retrying: {cell[:80]!r}")
            # Retry with low temperature, NO appended instructional text
            # translategemma echoes such instructions verbatim into the output
            # Just re-prompt at temp=0.1 with a tighter num_predict cap.
            try:
                prompt = self.translator._build_prompt(cell)
                retry_cap = max(20, len(cell) // 2 + 10)
                options = {"temperature": 0.1, "num_predict": retry_cap}
                if self.translator.seed is not None:
                    options["seed"] = self.translator.seed
                resp = ollama.chat(
                    model=self.translator.model,
                    messages=[{"role": "user", "content": prompt}],
                    options=options,
                )
                retry = resp["message"]["content"].strip()
            except Exception as e:
                _log_warning(f"  Retry failed: {e}")
                retry = ""
            if retry and not _looks_hallucinated(retry) and not _is_retry_echo(retry):
                return retry
            if _is_retry_echo(retry):
                _log_warning(f"  Retry echoed instruction text; keeping original: "
                             f"{cell[:80]!r}")
            else:
                _log_warning(f"  Retry also hallucinated; keeping original: "
                             f"{cell[:80]!r}")
            return cell
        return result

    def save_to_docx(self, translated_tables, referenced_indices, output_path):
        doc = Document()
        for idx in sorted(referenced_indices):
            if idx < len(translated_tables) and translated_tables[idx]:
                trans_table = translated_tables[idx]
                doc.add_paragraph(f"TABLE {idx + 1}")
                ncols = max(len(r) for r in trans_table)
                tbl = doc.add_table(rows=len(trans_table), cols=ncols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(trans_table):
                    for ci, cell in enumerate(row):
                        if ci < ncols:
                            tbl.rows[ri].cells[ci].text = cell
                doc.add_paragraph()  # spacing between tables
        doc.save(output_path)
        print(f"Saved translated tables to {output_path}")


class FootnoteTranslator:
    """Translates Footnote blocks and renders them as a docx table."""

    def __init__(self, translator):
        self.translator = translator

    def translate(self, footnotes):
        """Translate a list of Footnote blocks. 
        Returns dicts with number/text/translation, ready for save_to_docx."""
        print(f"  Translating {len(footnotes)} footnote(s)...")
        out = []
        for fn in footnotes:
            out.append({
                "number": fn.marker,
                "text": fn.text,
                "translation": self.translator.translate(fn.text),
            })
        return out

    def save_to_docx(self, footnotes, output_path):
        doc = Document()
        doc.add_paragraph("Translated Footnotes")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for i, h in enumerate(["#", "Original", "Translation"]):
            table.rows[0].cells[i].text = h

        for fn in footnotes:
            row = table.add_row()
            row.cells[0].text = fn["number"]
            row.cells[1].text = fn["text"]
            row.cells[2].text = fn["translation"]

        doc.save(output_path)
        print(f"Saved translated footnotes to {output_path}")


class CommentTranslator:
    """Translates Comment blocks and renders them as a docx table."""

    def __init__(self, translator):
        self.translator = translator

    def translate(self, comments):
        """Translate a list of Comment blocks. 
        Returns dicts with the original fields plus translation/anchor_translation."""
        out = []
        for i, c in enumerate(comments):
            print(f"  Translating comment {i+1}/{len(comments)}...")
            translation = self.translator.translate(c.text)
            anchor_translation = (self.translator.translate(c.anchor)
                                  if c.anchor else "")
            out.append({
                "id": c.id,
                "author": c.author,
                "date": c.date,
                "text": c.text,
                "anchor": c.anchor,
                "translation": translation,
                "anchor_translation": anchor_translation,
            })
        return out

    def save_to_docx(self, translated, output_path):
        doc = Document()
        doc.add_paragraph("Translated Comments")

        if not translated:
            doc.add_paragraph("No comments found.")
            doc.save(output_path)
            return

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Author", "Date", "Anchor (src)", "Comment (src)", "Translation"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for c in translated:
            row = table.add_row()
            row.cells[0].text = c["author"]
            row.cells[1].text = c["date"]
            row.cells[2].text = c["anchor"]
            row.cells[3].text = c["text"]
            row.cells[4].text = c["translation"]

        doc.save(output_path)
        print(f"Saved translated comments to {output_path}")


class DocumentTranslator:
    """Translates a list of Blocks and renders them to .docx.

    Used by both the PDF flow (PdfExtractor) and the docx flow (DocxExtractor).
    Tables and footnotes are routed to TableTranslator / FootnoteTranslator for their own output files.
    """

    def __init__(self, translator):
        self.translator = translator

    # ---- translation ---------------------------------------------------

    @staticmethod
    def _chunk_body_paras(body_indices, blocks, max_chars=1500, max_blocks=3):
        chunks: list[list[int]] = [[]]
        current_len = 0
        for i in body_indices:
            text_len = len(blocks[i].text)
            if chunks[-1] and (current_len + text_len + 1 > max_chars
                               or len(chunks[-1]) >= max_blocks):
                chunks.append([])
                current_len = 0
            chunks[-1].append(i)
            current_len += text_len + 1
        if not chunks[-1]:
            chunks.pop()
        return chunks

    def _translate_chunk_indices(self, indices, blocks):
        if len(indices) == 1:
            i = indices[0]
            return {i: self.translator.translate(blocks[i].to_markdown())}

        lines = [f"[{n}] {blocks[i].to_markdown()}" for n, i in enumerate(indices)]
        raw = self.translator.translate("\n".join(lines))

        parsed: dict[int, str] = {}
        current_n = None
        current_parts: list[str] = []
        for line in raw.split("\n"):
            m = _BLOCK_IDX_RE.match(line.strip())
            if m:
                n = int(m.group(1))
                if 0 <= n < len(indices):
                    if current_n is not None:
                        parsed[current_n] = " ".join(current_parts).strip()
                    current_n = n
                    current_parts = [line.strip()[m.end():].strip()]
                    continue
            if current_n is not None:
                current_parts.append(line.strip())
        if current_n is not None:
            parsed[current_n] = " ".join(current_parts).strip()

        # Length-ratio sanity check; outliers retry as singletons.
        results: dict[int, str] = {}
        for n, i in enumerate(indices):
            src_md = blocks[i].to_markdown()
            cand = parsed.get(n, "")
            ratio = len(cand) / max(len(src_md), 1)
            if cand and 0.4 <= ratio <= 2.5:
                results[i] = cand
            else:
                print(f"      Re-translating block {i} individually "
                      f"(chunk alignment/length mismatch: {len(cand)}/{len(src_md)})")
                results[i] = self.translator.translate(src_md)
        return results

    def _translate_blocks(self, blocks):
        results: dict = {}

        for i, b in enumerate(blocks):
            if isinstance(b, Heading):
                # Strip leading number prefix (e.g. "1.") before translating
                # so the LLM doesn't drop it, then re-prepend after.
                prefix = ""
                text = b.to_markdown()
                m = re.match(r'^(\d+(?:\.\d+)*[\.\)]\s*)', text)
                if m:
                    prefix = m.group(1)
                    text = text[m.end():]
                translated = _translate_preserving_fn(
                    text, self.translator.translate)
                results[i] = prefix + translated

        li_blocks = [(i, b) for i, b in enumerate(blocks)
                     if isinstance(b, ListItem)]
        if li_blocks:
            print(f"  Translating {len(li_blocks)} list-item(s)...")
            for i, b in li_blocks:
                title = self.translator.translate(b.title)
                m = _LIST_MARKER_RE.match(b.title)
                if m and not _LIST_MARKER_RE.match(title):
                    title = f"{m.group(1)} {title.lstrip()}"
                body = _translate_preserving_fn(
                    b.body_to_markdown(), self.translator.translate)
                results[i] = (title, body)

        body_idx = [i for i, b in enumerate(blocks) if isinstance(b, BodyPara)]
        # Use the markdown-encoded text for FN detection and translation.
        fn_idx = [i for i in body_idx
                  if _FN_MARKER_RE.search(blocks[i].to_markdown())]
        plain_idx = [i for i in body_idx if i not in set(fn_idx)]

        if fn_idx:
            print(f"  Translating {len(fn_idx)} paragraph(s) with footnotes...")
            for i in fn_idx:
                results[i] = _translate_preserving_fn(
                    blocks[i].to_markdown(), self.translator.translate)

        chunks = self._chunk_body_paras(plain_idx, blocks)
        print(f"  Translating {len(chunks)} body chunk(s) "
              f"({len(plain_idx)} paragraphs)...")
        for ci, chunk in enumerate(chunks):
            chars = sum(len(blocks[i].to_markdown()) for i in chunk)
            print(f"    Chunk {ci+1}/{len(chunks)} "
                  f"({len(chunk)} block(s), {chars} chars)...")
            results.update(self._translate_chunk_indices(chunk, blocks))

        return results

    # ---- rendering -----------------------------------------------------

    def _render_heading(self, doc, block: Heading, translated: str):
        heading = doc.add_heading("", level=min(block.level, 4))
        self._emit_text(heading, translated, size=None)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    @staticmethod
    def _emit_text(para, text: str, *, size: int | None = 11):
        """Emit text into a paragraph, parsing markdown emphasis and ‹FN› refs.

        Handles ***bold-italic***, **bold**, *italic*, and ‹FNn› superscripts.
        A ‹FNn› marker inside an emphasis span (e.g. `*Bedoya Lima v.‹FN3› Colombia*`) still becomes its own plain superscript run — footnote refs are annotations, not part of the surrounding formatting — while the text on either side keeps its italic/bold.
        """
        def add_run(t, *, italic=False, bold=False, superscript=False):
            if not t:
                return
            run = para.add_run(t)
            if size is not None:
                run.font.size = Pt(size)
            if italic:
                run.italic = True
            if bold:
                run.bold = True
            if superscript:
                run.font.superscript = True

        def emit_with_fn(t, *, italic=False, bold=False):
            # Split text on ‹FNn› markers so each marker becomes its own superscript run.
            # FN markers are footnote annotations, not part of the surrounding content
            # always rendered as plain superscript regardless of any italic/bold context.
            last_inner = 0
            for fm in _FN_MARKER_RE.finditer(t):
                add_run(t[last_inner:fm.start()], italic=italic, bold=bold)
                add_run(fm.group(1), superscript=True)
                last_inner = fm.end()
            add_run(t[last_inner:], italic=italic, bold=bold)

        last = 0
        for m in _SPAN_RE.finditer(text):
            if m.start() > last:
                emit_with_fn(text[last:m.start()])
            g1, g2, g3, g4 = m.group(1), m.group(2), m.group(3), m.group(4)
            if g4 is not None:          # top-level ‹FNn› is superscript
                add_run(g4, superscript=True)
            else:
                content = g1 or g2 or g3
                italic = bool(g1 or g3)
                bold = bool(g1 or g2)
                emit_with_fn(content, italic=italic, bold=bold)
            last = m.end()
        if last < len(text):
            emit_with_fn(text[last:])

    def _render_body(self, doc, block: BodyPara, translated: str):
        para = doc.add_paragraph()
        self._emit_text(para, translated)

    def _render_list_item(self, doc, block: ListItem,
                          trans_title: str, trans_body: str):
        para = doc.add_paragraph()
        self._emit_text(para, trans_title, size=11)
        sep = para.add_run(block.separator)
        sep.font.size = Pt(11)
        self._emit_text(para, trans_body)

    def _render_image(self, doc, block: ImageBlock):
        try:
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(io.BytesIO(block.data),
                            width=Inches(block.width_inches))
            return True
        except Exception as e:
            print(f"  Warning: could not insert image: {e}")
            return False

    # ---- orchestration -------------------------------------------------

    def translate_to_docx(self, blocks, tables, output_path):
        out_p = Path(output_path)
        tables_path = str(out_p.parent / f"{out_p.stem}_tables{out_p.suffix}")
        footnotes_path = str(out_p.parent / f"{out_p.stem}_footnotes{out_p.suffix}")
        comments_path = str(out_p.parent / f"{out_p.stem}_comments{out_p.suffix}")

        footnote_blocks = [b for b in blocks if isinstance(b, Footnote)]
        comment_blocks = [b for b in blocks if isinstance(b, Comment)]
        image_blocks = [b for b in blocks if isinstance(b, ImageBlock)]
        table_blocks = [b for b in blocks if isinstance(b, TablePlaceholder)]
        chars_in = sum(len(b.text) for b in blocks if isinstance(b, BodyPara))
        chars_in += sum(len(b.body_text) + len(b.title)
                        for b in blocks if isinstance(b, ListItem))
        chars_in += sum(len(b.text) for b in blocks if isinstance(b, Heading))

        print(f"Extracted {len(blocks)} blocks "
              f"({chars_in} chars, {len(table_blocks)} table(s), "
              f"{len(image_blocks)} image(s), "
              f"{len(footnote_blocks)} footnote(s), "
              f"{len(comment_blocks)} comment(s))")

        translated = self._translate_blocks(blocks)

        doc = Document()
        img_count = 0
        for i, block in enumerate(blocks):
            if isinstance(block, ImageBlock):
                if self._render_image(doc, block):
                    img_count += 1
            elif isinstance(block, TablePlaceholder):
                doc.add_paragraph(f"[TABLE {block.index}]")
            elif isinstance(block, Separator):
                doc.add_paragraph(block.text)
            elif isinstance(block, Heading):
                self._render_heading(doc, block, translated[i])
            elif isinstance(block, ListItem):
                t_title, t_body = translated[i]
                self._render_list_item(doc, block, t_title, t_body)
            elif isinstance(block, BodyPara):
                self._render_body(doc, block, translated[i])
            # Footnote and Comment blocks are routed below to their own
            # sibling output files; not rendered into the body docx.
        doc.save(output_path)
        print(f"Saved translated text to {output_path}"
              + (f" ({img_count} image(s) preserved)" if img_count else ""))

        if footnote_blocks:
            ft = FootnoteTranslator(self.translator)
            ft.save_to_docx(ft.translate(footnote_blocks), footnotes_path)

        if comment_blocks:
            cx = CommentTranslator(self.translator)
            cx.save_to_docx(cx.translate(comment_blocks), comments_path)

        if tables:
            print(f"  Translating {len(tables)} table(s)...")
            tt = TableTranslator(self.translator)
            translated_tables = tt.translate(tables)
            tt.save_to_docx(translated_tables,
                            set(range(len(translated_tables))), tables_path)

        return {
            "text_output": output_path,
            "tables_output": tables_path if tables else None,
            "footnotes_output": footnotes_path if footnote_blocks else None,
            "comments_output": comments_path if comment_blocks else None,
            "total_blocks": len(blocks),
            "total_tables": len(tables),
            "total_images": len(image_blocks),
            "total_footnotes": len(footnote_blocks),
            "total_comments": len(comment_blocks),
            "chars_in": chars_in,
        }


# Names of the two pipeline phases. Pass any subset to `translate_document`'s `phases=` parameter to run just those phases.
PHASE_BUILD_GLOSSARY = "build_glossary"   # Phase 1 — review document, write glossary file
PHASE_TRANSLATE = "translate"             # Phase 2 — translate document using the glossary
ALL_PHASES = (PHASE_BUILD_GLOSSARY, PHASE_TRANSLATE)


def _extract_blocks(filepath):
    """Run the right extractor for the file type and return (blocks, tables)."""
    if filepath.lower().endswith(".pdf"):
        from pdf_extract import PdfExtractor
        extractor = PdfExtractor(filepath)
        try:
            return extractor.extract()
        finally:
            extractor.close()
    from docx_extract import DocxExtractor
    extractor = DocxExtractor(filepath)
    try:
        return extractor.extract()
    finally:
        extractor.close()


def translate_document(filepath, output_path, *, source_lang="Spanish",
                       target_lang="English", model="translategemma",
                       review_model=None, glossary=None,
                       verbose_glossary=False, phases=ALL_PHASES,
                       force_rebuild=False, dump_dir=None, seed=42,
                       keep_glossary=True, timestamp=False,
                       domain=DEFAULT_DOMAIN):
    """Translate a .docx or .pdf file using the two-phase pipeline.

    Phases (pass as a tuple to `phases=`):
      "build_glossary" — Phase 1: review the document with the LLM, write a glossary file next to the output path. 
                         If a glossary file already exists at that path, it's loaded instead of regenerated (so your edits survive). 
                         Pass `force_rebuild=True` to delete and regenerate.
      "translate"      — Phase 2: translate the document using the glossary, render to .docx. 
                         Deletes the glossary file when done.

    Defaults run both phases.
    Pass `phases=("build_glossary",)` to stop after the glossary write (so you can edit it).
    Pass `phases=("translate",)` to skip review and reuse an existing glossary file — 
    this errors out if no glossary file exists yet, since otherwise the translation would silently run without any terminology rules.

    The `glossary` parameter controls glossary behavior:
      None           — default. Phase 1 builds one; Phase 2 uses it.
      False          — no glossary at all. Phase 1 is skipped; Phase 2 runs raw (no prompt injection, no violation checks).
      dict[str, str] — user-provided PREFER-only entries, merged into Phase 1.
      DomainGlossary — user-provided full glossary, merged into Phase 1.
      str | Path     — path to a glossary file; loaded as a user-provided glossary and merged into Phase 1.

    `force_rebuild=True` deletes any pre-existing glossary file before Phase 1 runs, forcing a fresh LLM review.
    Use this when you want to discard edits and re-derive the glossary from scratch.

    `dump_dir` (Phase 1 only) — when set, every per-segment input, prompt, raw LLM response, and parsed result from Step 1a/1b plus the merged state and final entries get written to that directory. 
    Phase 2 does not currently write any artifacts there. Useful for diffing prompt changes across runs.

    `seed` — passed to ollama as `options.seed` for every LLM call (Step 1a/1b, Phase 2 per-chunk translation, table cell retries).

    `keep_glossary` controls what happens to the glossary file after Phase 2 completes:
      True (default) — retain the file at its default location ({output}_glossary.txt).
                       Useful as a record of which terminology rules were applied.
      False          — delete the glossary file after Phase 2.
      str | Path     — copy the file to this path and delete the working copy.
                       Useful for dated archives, e.g.
                       `keep_glossary=f"archive/glossary_{date.today()}.txt"`.

    `timestamp` — when True, insert the current timestamp into the output stem (e.g. `foo.docx` -> `foo_2026-06-24_0915.docx`).
    # The glossary inherits the same timestamp, so each run produces a distinct (docx, glossary) pair that never overwrites prior runs. 

    Every translation completion appends a JSON line to `<output_dir>/translation_log.jsonl` recording timestamp, input, output, glossary, phases, and model.

    Returns the translation result dict, or `{"glossary_path": ...}` when only Phase 1 ran.
    """
    from entity_extract import DocumentReviewer

    unknown = set(phases) - set(ALL_PHASES)
    if unknown:
        raise ValueError(f"unknown phase(s) {sorted(unknown)}; expected from {ALL_PHASES}")
    if not phases:
        raise ValueError("phases must be non-empty")

    in_p = Path(filepath)
    if in_p.is_dir():
        raise ValueError(f"{filepath} is a folder. translate.py takes a single .pdf or .docx file; to translate every file in a folder, use batch_translate.py.")
    if not in_p.exists():
        raise ValueError(f"input file not found: {filepath}")
    if in_p.suffix.lower() not in (".pdf", ".docx"):
        raise ValueError(f"unsupported input type {in_p.suffix or '(no extension)'}: expected a .pdf or .docx file")

    blocks, tables = _extract_blocks(filepath)

    out_p = Path(output_path)
    run_ts = datetime.now().strftime("%Y-%m-%d_%H%M")
    if timestamp:
        # Insert timestamp before the extension so the glossary inherits it
        out_p = out_p.with_name(f"{out_p.stem}_{run_ts}{out_p.suffix}")
        output_path = str(out_p)
    out_p.parent.mkdir(parents=True, exist_ok=True)
    glossary_path = out_p.parent / f"{out_p.stem}_glossary.txt"

    coerced = DomainGlossary.coerce(glossary)
    no_glossary = coerced is False
    user_glossary = coerced if isinstance(coerced, DomainGlossary) else None

    if no_glossary and PHASE_BUILD_GLOSSARY in phases and PHASE_TRANSLATE not in phases:
        # User asked to build the glossary but also opted out of using one —
        # the only thing that makes sense to do is nothing.
        raise ValueError(
            "glossary=False with phases=('build_glossary',) is a no-op — "
            "either drop glossary=False to actually build, or drop the 'build_glossary' step."
        )

    resolved: DomainGlossary | None = None
    glossary_pre_existed = False

    if PHASE_BUILD_GLOSSARY in phases and not no_glossary:
        if force_rebuild and glossary_path.exists():
            print(f"  [translate_document] force_rebuild=True — removing existing glossary at {glossary_path}")
            DomainGlossary.delete(glossary_path)
        glossary_pre_existed = glossary_path.exists()
        resolved = DocumentReviewer(review_model or model, source_lang,
                                    target_lang, dump_dir=dump_dir,
                                    seed=seed).build_glossary(
            blocks, user_glossary, glossary_path)
    elif PHASE_TRANSLATE in phases and not no_glossary:
        # Translate-only + glossary expected → require an existing file
        if not glossary_path.exists():
            raise FileNotFoundError(
                f"phases=('translate',) requires an existing glossary at "
                f"{glossary_path}. Either:\n"
                f"  • Run with phases=('build_glossary',) first, or\n"
                f"  • Include 'build_glossary' in phases, or\n"
                f"  • Pass glossary=False to translate without one."
            )
        resolved = DomainGlossary.load(glossary_path, user_glossary=user_glossary)
    # else: no_glossary=True → resolved stays None

    if PHASE_BUILD_GLOSSARY in phases and PHASE_TRANSLATE not in phases:
        # Build-only path: glossary file stays on disk for the user to edit
        if glossary_pre_existed:
            print(f"\nExisting glossary at {glossary_path} reused as-is (pass force_rebuild=True to regenerate from the document).")
        else:
            print(f"\nGlossary written to {glossary_path}")
        print("Edit it if needed, then re-run with phases=('translate',) or the default both-phases to consume it.")
        _append_run_log(out_p.parent, {
            "timestamp": run_ts,
            "input": str(filepath),
            "output": None,
            "glossary": str(glossary_path),
            "phases": list(phases),
            "model": model,
            "regenerated": not glossary_pre_existed,
        })
        return {"glossary_path": str(glossary_path), "phases": list(phases), "regenerated": not glossary_pre_existed}

    if no_glossary:
        print("  [translate_document] glossary=False — Phase 2 will run without any glossary rules.")

    t = Translator(source_lang=source_lang, target_lang=target_lang,
                   model=model, glossary=resolved,
                   verbose_glossary=verbose_glossary, seed=seed, domain=domain)
    print(f"  [translate] persona/domain: {t.domain or '(general)'}. Phase 2 prompt template:")
    for line in t.prompt_preview().splitlines():
        print(f"    | {line}")
    dt = DocumentTranslator(t)
    result = dt.translate_to_docx(blocks, tables, output_path)
    result["input"] = filepath
    result["phases"] = list(phases)
    if keep_glossary is False:
        DomainGlossary.delete(glossary_path)
    elif keep_glossary is True:
        result["glossary_path"] = str(glossary_path)
        print(f"  [translate_document] glossary retained at {glossary_path}")
    else:
        archive = Path(keep_glossary)
        archive.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy(glossary_path, archive)
        DomainGlossary.delete(glossary_path)
        result["glossary_path"] = str(archive)
        print(f"  [translate_document] glossary archived to {archive}")
    _append_run_log(out_p.parent, {
        "timestamp": run_ts,
        "input": str(filepath),
        "output": str(out_p),
        "glossary": result.get("glossary_path"),  # None when keep_glossary=False
        "phases": list(phases),
        "model": model,
    })
    # Absolute-path summary of everything written
    saved = [result.get(k) for k in ("text_output", "tables_output", "footnotes_output", "comments_output") if result.get(k)]
    if result.get("glossary_path"):
        saved.append(result["glossary_path"])
    print(f"  [translate_document] saved {len(saved)} file(s):")
    for pth in saved:
        print(f"    {Path(pth).resolve()}")
    return result


def _append_run_log(output_dir: Path, entry: dict) -> None:
    """Append one JSON line to translation_log.jsonl in the output directory.
    Records timestamp + input/output/glossary paths + phases + model.
    """
    log_path = Path(output_dir) / "translation_log.jsonl"
    try:
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception as e:
        print(f"  [translate_document] could not append to {log_path}: {e}")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Translate a document (.docx or .pdf)")
    parser.add_argument("input_path",
                        help="Path to the input .docx or .pdf file")
    parser.add_argument("output_docx", 
                        help="Path for the translated output .docx")
    parser.add_argument("--source-lang", default="Spanish")
    parser.add_argument("--target-lang", default="English")
    parser.add_argument("--model", default="translategemma")
    parser.add_argument("--review-model", default=None,
                        help="Model for the glossary review pass (defaults to --model)")
    parser.add_argument("--glossary-only", action="store_true",
                        help="Shortcut for --phases build_glossary. Builds the glossary file and exits.")
    parser.add_argument("--translate-only", action="store_true",
                        help="Shortcut for --phases translate. Reuses an existing glossary file; errors if none exists at the expected path.")
    parser.add_argument("--phases", nargs="+", default=None,
                        choices=list(ALL_PHASES),
                        help=f"Pipeline phases to run (one or more of: {' '.join(ALL_PHASES)}). Default: all phases.")
    parser.add_argument("--force-rebuild", action="store_true",
                        help="Delete any pre-existing glossary file before the build phase, forcing a fresh LLM review.")
    parser.add_argument("--no-glossary", action="store_true",
                        help="Run the translation without any glossary rules. Skips prompt injection and violation checks.")
    parser.add_argument("--seed", type=int, default=42,
                        help="Integer seed for ollama generation (default 42).")
    parser.add_argument("--archive-glossary", metavar="PATH", default=None,
                        help="Copy the glossary file to PATH after Phase 2, then delete the working copy. Useful for dated archives, e.g. --archive-glossary archive/glossary_2026-06-21.txt.")
    parser.add_argument("--dump-dir", metavar="PATH", default=None,
                        help="Write Phase 1 debug snapshots (per-segment prompts, raw LLM responses, parsed results) to this directory.")
    parser.add_argument("--domain", default=DEFAULT_DOMAIN,
                        help='Subject-matter persona for the translation prompt (default: "human rights and public law"). Pass "" for a general translator.')
    parser.add_argument("--timestamp", action="store_true",
                        help="Insert the current timestamp into the output stem (e.g. foo_2026-06-24_0915.docx) so each run produces a distinct (docx, glossary) pair. Every run also appends a JSON line to translation_log.jsonl in the output directory.")
    args = parser.parse_args()

    if args.phases is not None:
        phases = tuple(args.phases)
    elif args.glossary_only:
        phases = (PHASE_BUILD_GLOSSARY,)
    elif args.translate_only:
        phases = (PHASE_TRANSLATE,)
    else:
        phases = ALL_PHASES

    try:
        translate_document(
            args.input_path, args.output_docx,
            source_lang=args.source_lang, target_lang=args.target_lang,
            model=args.model, review_model=args.review_model,
            phases=phases,
            force_rebuild=args.force_rebuild,
            glossary=False if args.no_glossary else None,
            seed=args.seed,
            keep_glossary=args.archive_glossary if args.archive_glossary else True,
            timestamp=args.timestamp,
            domain=args.domain,
            dump_dir=args.dump_dir,
        )
    except ValueError as e:
        raise SystemExit(f"error: {e}")
