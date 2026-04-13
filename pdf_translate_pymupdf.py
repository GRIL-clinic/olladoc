"""
PDF Translation using PyMuPDF
------------------------------
Extracts structured content from a PDF (text paragraphs, tables, images,
footnotes) and translates it to .docx, preserving heading hierarchy,
paragraph boundaries, per-run italic/bold formatting, list-item titles,
and footnote markers.

Pipeline:

    1. Spans: extract per-span text and formatting from PyMuPDF.
    2. Runs and Lines: collapse spans into formatted runs, grouped by line.
    3. Paragraphs: group adjacent lines into logical paragraphs using
       y-gaps and formatting transitions.
    4. Blocks: classify each paragraph as Heading, BodyPara, ListItem,
       Footnote, ImageBlock, TablePlaceholder, or Separator.
    5. Docx: translate and render each block by type.

Body text, tables, and footnotes are rendered to three .docx
files. Tables and footnotes are handled by TableTranslator and
FootnoteTranslator from translate.py; this module only classifies
them and hands off.

Usage:
    python pdf_translate_pymupdf.py input.pdf output.docx
"""

import fitz
import io
import re
from dataclasses import dataclass
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from pathlib import Path
from typing import List, Optional, Union

from translate import Translator, TableTranslator, FootnoteTranslator

fitz.no_recommend_layout = True


_BLOCK_IDX_RE = re.compile(r'^\[(\d+)\]\s*')
_ROMAN_HEADER_RE = re.compile(r'^[IVXLCDM]+\.\s*\S')
_ROMAN_ONLY_RE = re.compile(r'^[IVXLCDM]+\.$')
_LIST_ITEM_RE = re.compile(r'^\d+\)\s+\S')
_LIST_MARKER_RE = re.compile(r'^\s*(\d+[\).])\s*')
_LONE_MARKER_RE = re.compile(r'^(?:[*†‡§¶]+|\d+)$')
_SEPARATOR_RE = re.compile(r'^[_\-–—=]{3,}$')
_PAGE_NUMBER_RE = re.compile(r'^-?\s*\d+\s*-?$')
_DASH_SEPARATOR_RE = re.compile(r'^[\.\s]*[-–—][\.\s]*')


# ----- Data model ----------------------------------------------------------

@dataclass
class Run:
    text: str
    bold: bool = False
    italic: bool = False
    size: float = 0.0

    @property
    def stripped_len(self) -> int:
        return len(self.text.strip())


@dataclass
class Line:
    runs: List[Run]
    y0: float
    y1: float
    marker: Optional[str] = None

    @property
    def text(self) -> str:
        return "".join(r.text for r in self.runs).strip()


@dataclass
class Paragraph:
    runs: List[Run]
    y0: float
    y1: float
    marker: Optional[str] = None

    @property
    def text(self) -> str:
        return re.sub(r'\s+', ' ', "".join(r.text for r in self.runs)).strip()

    @property
    def dominant_size(self) -> float:
        sizes = {}
        for r in self.runs:
            c = r.stripped_len
            if c:
                sizes[r.size] = sizes.get(r.size, 0) + c
        return max(sizes, key=sizes.get) if sizes else 0.0

    @property
    def is_bold(self) -> bool:
        total = bold = 0
        for r in self.runs:
            c = r.stripped_len
            total += c
            if r.bold:
                bold += c
        return bool(total) and bold > total / 2

    @property
    def is_italic(self) -> bool:
        total = it = 0
        for r in self.runs:
            c = r.stripped_len
            total += c
            if r.italic:
                it += c
        return bool(total) and it > total / 2


@dataclass
class Heading:
    level: int
    runs: List[Run]

    @property
    def text(self) -> str:
        return re.sub(r'\s+', ' ', "".join(r.text for r in self.runs)).strip()


@dataclass
class BodyPara:
    runs: List[Run]

    @property
    def text(self) -> str:
        return re.sub(r'\s+', ' ', "".join(r.text for r in self.runs)).strip()


@dataclass
class ListItem:
    title: str
    body_runs: List[Run]
    separator: str = " "
    translated_title: Optional[str] = None

    @property
    def body_text(self) -> str:
        return re.sub(r'\s+', ' ', "".join(r.text for r in self.body_runs)).strip()


@dataclass
class Footnote:
    marker: str
    text: str


@dataclass
class ImageBlock:
    data: bytes
    width_inches: float


@dataclass
class TablePlaceholder:
    index: int


@dataclass
class Separator:
    text: str


Block = Union[Heading, BodyPara, ListItem, Footnote,
              ImageBlock, TablePlaceholder, Separator]


@dataclass
class PageContext:
    body_size: float
    body_is_bold: bool
    size_to_level: dict
    num_heading_sizes: int
    fn_threshold: float
    fn_separator_y: float


# ----- Extractor -----------------------------------------------------------

class PdfExtractor:
    """Parses a PDF into Blocks + raw tables. Footnotes are emitted inline
    as Footnote blocks; the caller routes them to FootnoteTranslator."""

    def __init__(self, filepath):
        self.doc = fitz.open(filepath)
        self.page_text_dicts = [
            page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            for page in self.doc
        ]
        (self.body_size, self.body_is_bold, self.size_to_level,
         self.num_heading_sizes, self.fn_threshold) = self._analyze_fonts()

    def close(self):
        self.doc.close()

    # ---- span helpers --------------------------------------------------

    @staticmethod
    def _has_font_trait(span, flag_bit, *keywords):
        if span.get("flags", 0) & (1 << flag_bit):
            return True
        font = span.get("font", "").lower()
        return any(kw in font for kw in keywords)

    @classmethod
    def _is_bold(cls, span):
        return cls._has_font_trait(span, 4, "bold", "heavy", "black")

    @classmethod
    def _is_italic(cls, span):
        return cls._has_font_trait(span, 1, "italic", "oblique")

    # ---- font analysis -------------------------------------------------

    def _analyze_fonts(self):
        size_chars, bold_chars = {}, {}
        for td in self.page_text_dicts:
            for block in td["blocks"]:
                if block["type"] != 0:
                    continue
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if not text:
                            continue
                        c = len(text)
                        sz = round(span["size"], 1)
                        size_chars[sz] = size_chars.get(sz, 0) + c
                        if self._is_bold(span):
                            bold_chars[sz] = bold_chars.get(sz, 0) + c
        if not size_chars:
            return None, False, {}, 0, 0.0
        body_size = max(size_chars, key=size_chars.get)
        body_total = size_chars.get(body_size, 0)
        body_bold = bold_chars.get(body_size, 0)
        body_is_bold = body_total > 0 and body_bold > body_total / 2
        larger = sorted((s for s in size_chars if s > body_size), reverse=True)
        size_to_level = {s: i + 1 for i, s in enumerate(larger[:4])}
        return (body_size, body_is_bold, size_to_level, len(larger),
                body_size * 0.85)

    # ---- Spans -> Lines ----------------------------------------

    @classmethod
    def _line_runs(cls, line_spans, first_line: bool):
        """Convert a PyMuPDF line's spans into Runs"""
        runs: List[Run] = []
        marker = None
        for si, span in enumerate(line_spans):
            text = span.get("text", "")
            if not text:
                continue
            if (first_line and si == 0 and marker is None
                    and len(text.strip()) <= 4 and len(line_spans) > 1):
                next_span = line_spans[si + 1]
                if span["size"] < next_span["size"] * 0.8:
                    marker = text.strip()
                    continue
            runs.append(Run(
                text=text,
                bold=cls._is_bold(span),
                italic=cls._is_italic(span),
                size=round(span["size"], 1),
            ))
        return runs, marker

    @classmethod
    def _extract_lines(cls, block) -> List[Line]:
        lines: List[Line] = []
        for li, line in enumerate(block["lines"]):
            runs, marker = cls._line_runs(line["spans"], first_line=(li == 0))
            if not any(r.stripped_len for r in runs):
                continue
            lines.append(Line(runs=runs, y0=line["bbox"][1],
                              y1=line["bbox"][3], marker=marker))
        return lines

    # ---- Lines -> Paragraphs -----------------------------------

    @staticmethod
    def _should_split(prev: Line, curr: Line, median_h: float) -> bool:
        gap = curr.y0 - prev.y1
        prev_size = max((r.size for r in prev.runs if r.stripped_len), default=0)
        curr_size = max((r.size for r in curr.runs if r.stripped_len), default=0)
        prev_bold = any(r.bold for r in prev.runs if r.stripped_len)
        curr_bold = any(r.bold for r in curr.runs if r.stripped_len)
        # Two consecutive bold lines are stacked titles, not wrapped text.
        both_bold = prev_bold and curr_bold
        return (
            gap > median_h * 0.7
            or abs(curr_size - prev_size) > 0.5
            or curr_bold != prev_bold
            or both_bold
            or (curr_bold and _ROMAN_HEADER_RE.match(curr.text))
            or _LIST_ITEM_RE.match(curr.text)
        )

    @staticmethod
    def _consolidate_runs(runs: List[Run]) -> List[Run]:
        out: List[Run] = []
        for r in runs:
            if out and out[-1].bold == r.bold and out[-1].italic == r.italic \
                    and out[-1].size == r.size:
                out[-1] = Run(text=out[-1].text + r.text, bold=r.bold,
                              italic=r.italic, size=r.size)
            else:
                out.append(r)
        return out

    @classmethod
    def _group_lines(cls, lines: List[Line]) -> List[Paragraph]:
        if not lines:
            return []
        heights = sorted(l.y1 - l.y0 for l in lines if l.y1 > l.y0)
        median_h = heights[len(heights) // 2] if heights else 10.0

        groups: List[List[Line]] = [[lines[0]]]
        for prev, curr in zip(lines, lines[1:]):
            if cls._should_split(prev, curr, median_h):
                groups.append([curr])
            else:
                groups[-1].append(curr)

        paragraphs: List[Paragraph] = []
        for group in groups:
            runs: List[Run] = []
            for i, line in enumerate(group):
                if i > 0 and runs:
                    last = runs[-1]
                    runs.append(Run(" ", bold=last.bold, italic=last.italic,
                                    size=last.size))
                runs.extend(line.runs)
            runs = cls._consolidate_runs(runs)
            if not any(r.stripped_len for r in runs):
                continue
            paragraphs.append(Paragraph(runs=runs, y0=group[0].y0,
                                        y1=group[-1].y1,
                                        marker=group[0].marker))
        return paragraphs

    # ---- Paragraphs -> Blocks ----------------------------------

    @staticmethod
    def _split_list_item(para: Paragraph):
        """If the paragraph has a leading italic title (with optional short
        non-italic "N)" preamble) followed by a non-italic body, return
        (title, body_runs, separator). Returns None if not a list item."""
        runs = para.runs
        if not runs:
            return None

        idx = 0
        preamble_parts: List[str] = []
        while idx < len(runs):
            r = runs[idx]
            if r.italic:
                break
            stripped = r.text.strip()
            if stripped and len(stripped) > 4:
                break
            preamble_parts.append(r.text)
            idx += 1

        italic_parts: List[str] = []
        while idx < len(runs) and runs[idx].italic:
            italic_parts.append(runs[idx].text)
            idx += 1

        if not italic_parts or idx == len(runs):
            return None

        title = re.sub(r'\s+', ' ',
                       "".join(preamble_parts) + "".join(italic_parts)).strip()
        body_runs = runs[idx:]
        body_text = "".join(r.text for r in body_runs).strip()

        if len(title) < 8 or len(body_text) < 10:
            return None

        # Title->body separator: ":" lifts onto title, dash separator
        # canonicalizes to " - "/" – "/" — ", otherwise plain space.
        separator = " "
        body_text_clean = body_text
        if body_text.startswith(":") and not title.endswith(":"):
            title = title.rstrip(" .") + ":"
            body_text_clean = body_text[1:].lstrip()
        else:
            m = _DASH_SEPARATOR_RE.match(body_text)
            if m:
                dash_match = re.search(r'[-–—]', m.group(0))
                dash = dash_match.group(0) if dash_match else '-'
                had_period = '.' in m.group(0)
                title = title.rstrip('. ')
                if had_period:
                    title = title + '.'
                separator = f" {dash} "
                body_text_clean = body_text[m.end():].lstrip()

        leading = next((r for r in body_runs if r.text.strip()), None)
        if leading is None:
            return None
        new_body = [Run(body_text_clean, bold=leading.bold,
                        italic=leading.italic, size=leading.size)]
        return title, new_body, separator

    @classmethod
    def _classify_paragraph(cls, para: Paragraph,
                            ctx: PageContext) -> Optional[Block]:
        text = para.text
        if not text:
            return None

        if _SEPARATOR_RE.match(text):
            return Separator(text=text)

        effective_bold = para.is_bold and not ctx.body_is_bold
        if (para.dominant_size < ctx.fn_threshold
                and para.y0 > ctx.fn_separator_y):
            if _PAGE_NUMBER_RE.match(text):
                return BodyPara(runs=list(para.runs))
            marker = para.marker
            body = text
            if not marker:
                m = re.match(r'^(\d+|[*†‡§¶])\s+', body)
                if m:
                    marker = m.group(1)
                    body = body[m.end():]
            return Footnote(marker=marker or "*", text=body)

        level = ctx.size_to_level.get(para.dominant_size, 0)
        is_roman = bool(_ROMAN_HEADER_RE.match(text)) and len(text) < 120
        if level == 0 and effective_bold and is_roman:
            level = min(ctx.num_heading_sizes + 1, 3)
        elif level == 0 and effective_bold and len(text) < 120:
            level = min(ctx.num_heading_sizes + 1, 4)
        if level > 0:
            return Heading(level=level, runs=list(para.runs))

        li = cls._split_list_item(para)
        if li is not None:
            title, body_runs, separator = li
            return ListItem(title=title, body_runs=body_runs,
                            separator=separator)

        return BodyPara(runs=list(para.runs))

    # ---- post-classification cleanup -----------------------------------

    @staticmethod
    def _fold_lone_markers(blocks: List[Block]) -> List[Block]:
        """Fold a lone Roman-numeral Heading onto the next Heading; fold
        a stub Footnote onto the next Footnote."""
        out: List[Block] = []
        i = 0
        while i < len(blocks):
            cur = blocks[i]
            nxt = blocks[i + 1] if i + 1 < len(blocks) else None

            if (isinstance(cur, Heading) and _ROMAN_ONLY_RE.match(cur.text)
                    and isinstance(nxt, Heading)):
                prefix = cur.text + " "
                new_runs = list(nxt.runs)
                for j, r in enumerate(new_runs):
                    if r.stripped_len:
                        new_runs[j] = Run(prefix + r.text, bold=r.bold,
                                          italic=r.italic, size=r.size)
                        break
                out.append(Heading(level=nxt.level, runs=new_runs))
                i += 2
                continue

            if (isinstance(cur, Footnote)
                    and _LONE_MARKER_RE.match(cur.text.strip())
                    and isinstance(nxt, Footnote)):
                marker = cur.marker if cur.marker and cur.marker != "*" \
                    else (cur.text.strip() or nxt.marker or "*")
                out.append(Footnote(marker=marker, text=nxt.text))
                i += 2
                continue

            out.append(cur)
            i += 1
        return out

    @staticmethod
    def _is_continuation(prev_text: str, next_text: str) -> bool:
        prev_tail = prev_text.rstrip()
        if not prev_tail or prev_tail[-1] in ".!?:;»)”\"":
            return False
        if not next_text:
            return False
        first = next_text[0]
        return first.islower() or first in "(,;"

    @classmethod
    def _merge_body_continuations(cls, blocks: List[Block]) -> List[Block]:
        """Stitch paragraphs split across PyMuPDF block boundaries / page
        breaks. Extends BodyPara or ListItem.body_runs."""
        out: List[Block] = []
        for b in blocks:
            if isinstance(b, ListItem):
                out.append(b)
                continue
            if not out or not isinstance(b, BodyPara):
                out.append(b)
                continue
            prev = out[-1]
            if isinstance(prev, BodyPara) and cls._is_continuation(prev.text, b.text):
                joiner = Run(" ", size=prev.runs[-1].size if prev.runs else 0.0)
                out[-1] = BodyPara(runs=cls._consolidate_runs(
                    prev.runs + [joiner] + list(b.runs)))
                continue
            if isinstance(prev, ListItem) and cls._is_continuation(prev.body_text, b.text):
                joiner = Run(
                    " ", size=prev.body_runs[-1].size if prev.body_runs else 0.0)
                out[-1] = ListItem(
                    title=prev.title,
                    body_runs=cls._consolidate_runs(
                        prev.body_runs + [joiner] + list(b.runs)),
                    separator=prev.separator,
                )
                continue
            out.append(b)
        return out

    # ---- table / image extraction --------------------------------------

    @staticmethod
    def _render_rect(page, rect, scale=2.0):
        pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), clip=rect)
        return pix.tobytes("png")

    @staticmethod
    def _collapse_merged_columns(raw_rows):
        """Collapse PyMuPDF's raw grid for a table with merged cells."""
        if not raw_rows:
            return raw_rows
        ncols = max(len(r) for r in raw_rows)
        if ncols <= 1:
            return raw_rows
        merged_rows = []
        for row in raw_rows:
            padded = [(row[i] if i < len(row) else "") or "" for i in range(ncols)]
            collapsed = [padded[0]]
            for ci in range(1, ncols):
                prev_c = collapsed[-1].strip()
                curr_c = padded[ci].strip()
                if not prev_c and not curr_c:
                    collapsed[-1] = ""
                elif not prev_c:
                    collapsed[-1] = curr_c
                elif not curr_c:
                    pass
                else:
                    collapsed.append(curr_c)
            merged_rows.append(collapsed)
        max_cols = max(len(r) for r in merged_rows)
        for row in merged_rows:
            while len(row) < max_cols:
                row.append("")
        keep = [ci for ci in range(max_cols)
                if any(merged_rows[ri][ci].strip() for ri in range(len(merged_rows)))]
        if len(keep) == max_cols:
            return merged_rows
        return [[row[ci] for ci in keep] for row in merged_rows]

    def _find_image_regions(self, page, table_rects, drawings, min_size=50):
        expanded = [fitz.Rect(r.x0 - 5, r.y0 - 5, r.x1 + 5, r.y1 + 5)
                    for r in table_rects]
        results = []
        seen = set()
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            for rect in page.get_image_rects(xref):
                if rect.width < min_size or rect.height < min_size:
                    continue
                if any(rect.intersects(tr) for tr in expanded):
                    continue
                key = (round(rect.x0), round(rect.y0),
                       round(rect.x1), round(rect.y1))
                if key in seen:
                    continue
                seen.add(key)
                png = self._render_rect(page, rect)
                results.append((fitz.Rect(rect), png, min(rect.width / 72.0, 6.0)))
        draw_rects = []
        for d in drawings:
            r = fitz.Rect(d["rect"])
            if abs(r.height) < 3 and r.width > 30:
                continue
            if r.width < 3 and r.height < 3:
                continue
            if any(r.intersects(tr) for tr in expanded):
                continue
            draw_rects.append(r)
        if len(draw_rects) >= 10:
            union = fitz.Rect(draw_rects[0])
            for r in draw_rects[1:]:
                union |= r
            if (union.width >= min_size and union.height >= min_size
                    and not any(union.intersects(fitz.Rect(*k)) for k in seen)):
                png = self._render_rect(page, union)
                results.append((union, png, min(union.width / 72.0, 6.0)))
        return results

    @staticmethod
    def _find_footnote_separator_y(drawings, page_height):
        for d in drawings:
            rect = d["rect"]
            if (abs(rect.height) < 3 and 30 < rect.width < 300
                    and rect.y0 > page_height * 0.25):
                return rect.y0
        return page_height * 0.6

    # ---- per-page pipeline ---------------------------------------------

    def _process_page(self, page, text_dict, table_idx):
        page_height = page.rect.height
        drawings = page.get_drawings()
        ctx = PageContext(
            body_size=self.body_size,
            body_is_bold=self.body_is_bold,
            size_to_level=self.size_to_level,
            num_heading_sizes=self.num_heading_sizes,
            fn_threshold=self.fn_threshold,
            fn_separator_y=self._find_footnote_separator_y(
                drawings, page_height),
        )

        page_table_info = []
        for table in page.find_tables().tables:
            raw_rows = [[cell or "" for cell in row] for row in table.extract()]
            if raw_rows:
                page_table_info.append((
                    fitz.Rect(table.bbox),
                    self._collapse_merged_columns(raw_rows),
                ))
        table_bboxes = [info[0] for info in page_table_info]

        image_info = self._find_image_regions(page, table_bboxes, drawings)

        page_elements: list[tuple[float, Block]] = []
        page_tables = []
        for tbbox, tdata in page_table_info:
            page_tables.append(tdata)
            table_idx += 1
            page_elements.append((tbbox.y0, TablePlaceholder(index=table_idx)))
        for rect, png, w in image_info:
            page_elements.append((rect.y0,
                                  ImageBlock(data=png, width_inches=w)))

        for block in text_dict["blocks"]:
            if block["type"] != 0:
                continue
            block_rect = fitz.Rect(block["bbox"])
            if any(block_rect.intersects(r) for r, _ in page_table_info):
                continue
            if any(block_rect.intersects(rect) for rect, *_ in image_info):
                continue
            lines = self._extract_lines(block)
            for para in self._group_lines(lines):
                blk = self._classify_paragraph(para, ctx)
                if blk is not None:
                    page_elements.append((para.y0, blk))

        page_elements.sort(key=lambda x: x[0])
        return [b for _, b in page_elements], page_tables, table_idx

    # ---- public API ----------------------------------------------------

    def extract(self) -> tuple[List[Block], list]:
        if self.body_size is None:
            return [], []
        all_blocks: List[Block] = []
        all_tables = []
        table_idx = 0
        for page, td in zip(self.doc, self.page_text_dicts):
            page_blocks, page_tables, table_idx = self._process_page(
                page, td, table_idx)
            all_blocks.extend(page_blocks)
            all_tables.extend(page_tables)
        all_blocks = self._fold_lone_markers(all_blocks)
        all_blocks = self._merge_body_continuations(all_blocks)
        return all_blocks, all_tables


# ----- Translator ----------------------------------------------------------

class PdfDocumentTranslator:
    """Translates a list of Blocks and renders them to .docx."""

    def __init__(self, translator):
        self.translator = translator

    # ---- translation ---------------------------------------------------

    @staticmethod
    def _chunk_body_paras(body_indices, blocks, max_chars=1500, max_blocks=3):
        chunks: list[list[int]] = [[]]
        current_len = 0
        for i in body_indices:
            text_len = len(blocks[i].text)
            if chunks[-1] and (current_len + text_len + 1 > max_chars
                               or len(chunks[-1]) >= max_blocks):
                chunks.append([])
                current_len = 0
            chunks[-1].append(i)
            current_len += text_len + 1
        if not chunks[-1]:
            chunks.pop()
        return chunks

    def _translate_chunk_indices(self, indices, blocks):
        if len(indices) == 1:
            i = indices[0]
            return {i: self.translator.translate(blocks[i].text)}

        lines = [f"[{n}] {blocks[i].text}" for n, i in enumerate(indices)]
        raw = self.translator.translate("\n".join(lines))

        parsed: dict[int, str] = {}
        current_n = None
        current_parts: list[str] = []
        for line in raw.split("\n"):
            m = _BLOCK_IDX_RE.match(line.strip())
            if m:
                n = int(m.group(1))
                if 0 <= n < len(indices):
                    if current_n is not None:
                        parsed[current_n] = " ".join(current_parts).strip()
                    current_n = n
                    current_parts = [line.strip()[m.end():].strip()]
                    continue
            if current_n is not None:
                current_parts.append(line.strip())
        if current_n is not None:
            parsed[current_n] = " ".join(current_parts).strip()

        results: dict[int, str] = {}
        for n, i in enumerate(indices):
            src_text = blocks[i].text
            cand = parsed.get(n, "")
            ratio = len(cand) / max(len(src_text), 1)
            if cand and 0.4 <= ratio <= 2.5:
                results[i] = cand
            else:
                print(f"      Re-translating block {i} individually "
                      f"(chunk alignment/length mismatch: {len(cand)}/{len(src_text)})")
                results[i] = self.translator.translate(src_text)
        return results

    def _translate_blocks(self, blocks: List[Block]) -> dict:
        results: dict = {}

        for i, b in enumerate(blocks):
            if isinstance(b, Heading):
                results[i] = self.translator.translate(b.text)

        li_blocks = [(i, b) for i, b in enumerate(blocks)
                     if isinstance(b, ListItem)]
        if li_blocks:
            print(f"  Translating {len(li_blocks)} list-item(s)...")
            for i, b in li_blocks:
                title = self.translator.translate(b.title)
                # Re-prepend a leading "N)" if Gemma dropped it.
                m = _LIST_MARKER_RE.match(b.title)
                if m and not _LIST_MARKER_RE.match(title):
                    title = f"{m.group(1)} {title.lstrip()}"
                body = self.translator.translate(b.body_text)
                results[i] = (title, body)

        body_idx = [i for i, b in enumerate(blocks) if isinstance(b, BodyPara)]
        chunks = self._chunk_body_paras(body_idx, blocks)
        print(f"  Translating {len(chunks)} body chunk(s) "
              f"({len(body_idx)} paragraphs)...")
        for ci, chunk in enumerate(chunks):
            chars = sum(len(blocks[i].text) for i in chunk)
            print(f"    Chunk {ci+1}/{len(chunks)} "
                  f"({len(chunk)} block(s), {chars} chars)...")
            results.update(self._translate_chunk_indices(chunk, blocks))

        return results

    # ---- rendering -----------------------------------------------------

    def _render_heading(self, doc, block: Heading, translated: str):
        heading = doc.add_heading(translated, level=min(block.level, 4))
        any_italic = any(r.italic for r in block.runs if r.stripped_len)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            if any_italic:
                run.italic = True

    def _render_body(self, doc, block: BodyPara, translated: str):
        para = doc.add_paragraph()
        run = para.add_run(translated)
        run.font.size = Pt(11)
        strip_runs = [r for r in block.runs if r.stripped_len]
        if strip_runs and all(r.bold for r in strip_runs):
            run.bold = True
        if strip_runs and all(r.italic for r in strip_runs):
            run.italic = True

    def _render_list_item(self, doc, block: ListItem,
                          trans_title: str, trans_body: str):
        para = doc.add_paragraph()
        t = para.add_run(trans_title)
        t.font.size = Pt(11)
        t.italic = True
        sep = para.add_run(block.separator)
        sep.font.size = Pt(11)
        b = para.add_run(trans_body)
        b.font.size = Pt(11)

    def _render_image(self, doc, block: ImageBlock):
        try:
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(io.BytesIO(block.data),
                            width=Inches(block.width_inches))
            return True
        except Exception as e:
            print(f"  Warning: could not insert image: {e}")
            return False

    # ---- orchestration -------------------------------------------------

    def translate_to_docx(self, blocks: List[Block], tables: list,
                          output_path: str) -> dict:
        out_p = Path(output_path)
        tables_path = str(out_p.parent / f"{out_p.stem}_tables{out_p.suffix}")
        footnotes_path = str(out_p.parent / f"{out_p.stem}_footnotes{out_p.suffix}")

        footnote_blocks = [b for b in blocks if isinstance(b, Footnote)]
        image_blocks = [b for b in blocks if isinstance(b, ImageBlock)]
        table_blocks = [b for b in blocks if isinstance(b, TablePlaceholder)]
        chars_in = sum(len(b.text) for b in blocks if isinstance(b, BodyPara))
        chars_in += sum(len(b.body_text) + len(b.title)
                        for b in blocks if isinstance(b, ListItem))
        chars_in += sum(len(b.text) for b in blocks if isinstance(b, Heading))

        print(f"Extracted {len(blocks)} blocks "
              f"({chars_in} chars, {len(table_blocks)} table(s), "
              f"{len(image_blocks)} image(s), "
              f"{len(footnote_blocks)} footnote(s))")

        translated = self._translate_blocks(blocks)

        doc = Document()
        img_count = 0
        for i, block in enumerate(blocks):
            if isinstance(block, ImageBlock):
                if self._render_image(doc, block):
                    img_count += 1
            elif isinstance(block, TablePlaceholder):
                doc.add_paragraph(f"[TABLE {block.index}]")
            elif isinstance(block, Separator):
                doc.add_paragraph(block.text)
            elif isinstance(block, Heading):
                self._render_heading(doc, block, translated[i])
            elif isinstance(block, ListItem):
                t_title, t_body = translated[i]
                self._render_list_item(doc, block, t_title, t_body)
            elif isinstance(block, BodyPara):
                self._render_body(doc, block, translated[i])
        doc.save(output_path)
        print(f"Saved translated text to {output_path}"
              + (f" ({img_count} image(s) preserved)" if img_count else ""))

        if footnote_blocks:
            fn_dicts = [{"number": f.marker, "text": f.text}
                        for f in footnote_blocks]
            ft = FootnoteTranslator(self.translator)
            ft.translate(fn_dicts)
            ft.save_to_docx(fn_dicts, footnotes_path)

        if tables:
            print(f"  Translating {len(tables)} table(s)...")
            tt = TableTranslator(self.translator)
            translated_tables = tt.translate(tables)
            tt.save_to_docx(translated_tables,
                            set(range(len(translated_tables))), tables_path)

        return {
            "text_output": output_path,
            "tables_output": tables_path if tables else None,
            "footnotes_output": footnotes_path if footnote_blocks else None,
            "total_blocks": len(blocks),
            "total_tables": len(tables),
            "total_images": len(image_blocks),
            "total_footnotes": len(footnote_blocks),
            "chars_in": chars_in,
        }


def translate_pdf(filepath, output_path, source_lang="Spanish",
                  target_lang="English", model="translategemma",
                  glossary=None):
    extractor = PdfExtractor(filepath)
    try:
        blocks, tables = extractor.extract()
    finally:
        extractor.close()

    t = Translator(source_lang=source_lang, target_lang=target_lang,
                   model=model, glossary=glossary)
    dt = PdfDocumentTranslator(t)
    return dt.translate_to_docx(blocks, tables, output_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Translate a PDF with formatting (PyMuPDF)")
    parser.add_argument("input_pdf", help="Path to the input PDF file")
    parser.add_argument("output_docx", help="Path for the translated output .docx")
    parser.add_argument("--source-lang", default="Spanish")
    parser.add_argument("--target-lang", default="English")
    parser.add_argument("--model", default="translategemma")
    args = parser.parse_args()

    translate_pdf(
        args.input_pdf, args.output_docx,
        source_lang=args.source_lang, target_lang=args.target_lang,
        model=args.model,
    )
